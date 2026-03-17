"""
documents.py — Phase 2 document generation
Produces a tailored CV and cover letter as .docx files for every job that
scores 75+ and has proceed=true.
"""

import json
import re
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import PROFILE_PATH, BASE_DIR
from get_brand_color import get_brand_color

OUTPUTS_DIR = BASE_DIR / "outputs"

# Load profile once
_PROFILE: dict | None = None


def _profile() -> dict:
    global _PROFILE
    if _PROFILE is None:
        _PROFILE = json.loads(PROFILE_PATH.read_text())
    return _PROFILE


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_documents(job: dict) -> tuple[str, str]:
    """
    Generate a tailored CV and cover letter for a job that passed scoring.
    Returns (cv_path, cover_letter_path) as strings relative to BASE_DIR.
    """
    OUTPUTS_DIR.mkdir(exist_ok=True)
    profile = _profile()

    slug = _make_slug(job)
    today = date.today().strftime("%Y-%m-%d")

    cv_path = OUTPUTS_DIR / f"cv_{slug}_{today}.docx"
    cl_path = OUTPUTS_DIR / f"coverletter_{slug}_{today}.docx"

    matched_skills = json.loads(job.get("matched_skills") or "[]")

    _build_cv(profile, job, matched_skills, cv_path)
    _build_cover_letter(profile, job, cl_path)

    # Return paths relative to project root for clean display
    return str(cv_path.relative_to(BASE_DIR)), str(cl_path.relative_to(BASE_DIR))


_ROLE_WORD = re.compile(
    r"\b(Lead|Director|Manager|Officer|Head|Partner|Advisor|Associate|Consultant|Specialist|Executive)\b"
)


def _clean_job_title(title: str) -> str:
    """Insert ' / ' between two concatenated role titles if detected.

    e.g. 'Technology, Product and Innovation Lead Incubation Lead'
      -> 'Technology, Product and Innovation Lead / Incubation Lead'
    """
    matches = list(_ROLE_WORD.finditer(title))
    if len(matches) >= 2:
        end = matches[0].end()
        if re.match(r"\s+[A-Z]", title[end:]):
            return title[:end] + " /" + title[end:]
    return title


def _make_slug(job: dict) -> str:
    def clean(s: str) -> str:
        s = re.sub(r"[^\w\s]", "", s or "unknown").strip()[:32]
        return re.sub(r"\s+", "_", s).lower()

    slug = f"{clean(job.get('company'))}_{clean(job.get('title'))}"
    return re.sub(r"_+", "_", slug).strip("_")


def _parse_tailored_cv(tailored_cv_text: str | None) -> dict | None:
    """Parse structured JSON from tailored_cv field.

    Returns a dict with 'roles' and optional 'projects' keys, or None if the
    field is absent, empty, or plain-text (old format — graceful fallback).
    """
    if not tailored_cv_text:
        return None
    try:
        text = tailored_cv_text.strip()
        # Strip markdown code fences Claude sometimes wraps JSON in
        if text.startswith("```"):
            lines = text.split("\n")
            text = "\n".join(lines[1:-1])
        data = json.loads(text)
        if isinstance(data, dict) and "roles" in data:
            return data
    except (json.JSONDecodeError, ValueError):
        pass
    return None


# ---------------------------------------------------------------------------
# Shared style helpers
# ---------------------------------------------------------------------------

def _font(run, size: float, bold=False, italic=False, color: tuple | None = None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _spacing(para, before=0, after=2, line=None):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line:
        fmt.line_spacing = Pt(line)


def _section_heading(doc: Document, text: str):
    """11pt bold Arial, no border, 14pt before / 3pt after — matches master_cv.docx."""
    p = doc.add_paragraph()
    _spacing(p, before=14, after=3)
    run = p.add_run(text.upper())
    _font(run, 11, bold=True)
    return p


def _remove_table_borders(tbl):
    """Strip all borders from a table (used for layout tables)."""
    tbl_el = tbl._tbl
    tbl_pr = tbl_el.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_el.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{name}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        tbl_borders.append(b)
    tbl_pr.append(tbl_borders)


# ---------------------------------------------------------------------------
# CV content extraction (produces a dict for the JS renderer)
# ---------------------------------------------------------------------------

def _extract_cv_content(profile: dict, job: dict, matched_skills: list[str]) -> dict:
    """Serialise all CV content into a plain dict consumed by generate_cv.js."""
    personal = profile["personal"]
    metrics  = profile.get("key_metrics", {})

    # Tagline — first two target roles + domain shorthand
    target_roles = profile.get("job_preferences", {}).get("target_roles", [])
    tagline_parts = list(dict.fromkeys(target_roles[:2])) + ["AI & Ventures"]
    tagline = "  ·  ".join(tagline_parts)

    contact = [
        personal.get("location", ""),
        personal.get("email", ""),
        personal.get("phone", ""),
        personal.get("linkedin", "").replace("https://www.", "").replace("https://", ""),
    ]
    contact = [c for c in contact if c]

    stats = [
        {"value": f"{metrics.get('years_experience', 23)} Years", "label": "Experience"},
        {"value": f"{metrics.get('startups_mentored', 200)}+",    "label": "Startups Mentored"},
        {"value": str(metrics.get("startups_incubated", 88)),     "label": "Startups Incubated"},
        {"value": str(metrics.get("portfolio_raised", "£12m+")), "label": "Portfolio Raised"},
        {"value": str(metrics.get("portfolio_market_cap_managed", "$8B+")), "label": "Portfolio Managed"},
    ]

    # Parse structured tailored_cv
    structured = _parse_tailored_cv(job.get("tailored_cv"))
    structured_bullets: dict[str, list[str]] = {}
    structured_projects: list[dict] = []
    tailored_summary: str | None = None
    if structured:
        tailored_summary = structured.get("summary") or None
        for role_data in structured.get("roles", []):
            key = role_data.get("company", "").lower()
            if key:
                structured_bullets[key] = role_data.get("bullets", [])
        structured_projects = structured.get("projects", [])

    # Projects — max 2, max 3 bullets each
    projects = [
        {"name": p.get("name", ""), "date": p.get("date", ""), "bullets": p.get("bullets", [])[:3]}
        for p in structured_projects[:2]
    ]

    # Split roles: main vs earlier career
    matched_lower = [s.lower() for s in matched_skills]
    main_roles, earlier_roles = [], []
    for role in profile["career_history"]:
        start_year = int(str(role["start"])[:4])
        is_old_parttime = role.get("type") == "part-time" and start_year < 2020
        if start_year < 2007 or is_old_parttime:
            earlier_roles.append(role)
        else:
            main_roles.append(role)

    experience = []
    for role in main_roles:
        start_year = int(str(role["start"])[:4])
        max_bullets = 4 if start_year >= 2015 else 2

        company_key = role["company"].lower()
        if company_key in structured_bullets:
            bullets = structured_bullets[company_key]
        else:
            bullets = list(role.get("achievements", []))
            bullets.sort(key=lambda b: -sum(1 for s in matched_lower if s in b.lower()))

        experience.append({
            "title":    role["title"],
            "company":  role["company"],
            "dates":    f"{role['start']} – {str(role['end']).capitalize()}",
            "location": role.get("location", ""),
            "bullets":  bullets[:max_bullets],
        })

    earlier_career = [
        {
            "company": r["company"],
            "title":   r["title"],
            "dates":   f"{r['start']} – {str(r['end']).capitalize()}",
        }
        for r in earlier_roles
    ]

    # Education — group AI certifications into one line, then key edu items
    education: list[str] = []
    certs = profile.get("ai_certifications", [])
    if certs:
        issuer = certs[0]["issuer"]
        year   = certs[0]["year"]
        titles = ", ".join(c["title"] for c in certs)
        education.append(f"{issuer} Certifications ({year}): {titles}")
    for edu in profile.get("education", [])[:3]:
        education.append(edu)

    # Skills table rows — use tailored override when present, else derive from profile
    if structured and structured.get("skills"):
        skills = structured["skills"]
    else:
        skills_data = profile.get("skills", {})
        SKILL_ROWS = [
            ("AI & Technology",      ["ai_specific", "technology"]),
            ("Programme Leadership", ["programme_leadership"]),
            ("Venture & Investment", ["venture_and_investment"]),
            ("Leadership",           ["leadership"]),
        ]
        skills = []
        for label, keys in SKILL_ROWS:
            merged: list[str] = []
            for k in keys:
                merged.extend(skills_data.get(k, []))
            skills.append({"label": label, "values": "  ·  ".join(merged[:5])})

    return {
        "name":          personal["name"],
        "tagline":       tagline,
        "contact":       contact,
        "stats":         stats,
        "summary":       tailored_summary or profile["summary"],
        "projects":      projects,
        "experience":    experience,
        "earlier_career": earlier_career,
        "education":     education,
        "skills":        skills,
    }


# ---------------------------------------------------------------------------
# CV builder — delegates rendering to generate_cv.js
# ---------------------------------------------------------------------------

def _render_cv_with_js(content: dict, accent: str, path: Path) -> None:
    """Pipe content JSON to generate_cv.js and write the styled .docx."""
    js_script    = BASE_DIR / "generate_cv.js"
    node_modules = BASE_DIR / "node_modules" / "docx"

    if not node_modules.exists():
        print("[Documents] Installing npm packages...")
        subprocess.run(["npm", "install"], cwd=BASE_DIR, check=True, capture_output=True)

    payload = json.dumps({"content": content, "accent": accent, "output": str(path)})
    result  = subprocess.run(
        ["node", str(js_script)],
        input=payload.encode(),
        capture_output=True,
        cwd=BASE_DIR,
    )
    if result.returncode != 0:
        raise RuntimeError(f"[Documents] generate_cv.js failed:\n{result.stderr.decode()}")
    if result.stdout:
        print(result.stdout.decode().strip())


def _build_cv(profile: dict, job: dict, matched_skills: list[str], path: Path):
    accent  = get_brand_color(job.get("company", ""))
    content = _extract_cv_content(profile, job, matched_skills)
    _render_cv_with_js(content, accent, path)
    print(f"[Documents] CV saved: {path.name}")


# ---------------------------------------------------------------------------
# Cover letter builder
# ---------------------------------------------------------------------------

def _build_cover_letter(profile: dict, job: dict, path: Path):
    doc = Document()

    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.15)
        sec.right_margin = Inches(1.15)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)

    personal = profile["personal"]

    def line(text: str, bold=False, italic=False, size=11,
             color: tuple | None = None, after=2):
        p = doc.add_paragraph()
        _spacing(p, before=0, after=after)
        _font(p.add_run(text), size, bold=bold, italic=italic, color=color)
        return p

    # ── Sender block ──────────────────────────────────────────────────────────
    line(personal["name"], bold=True, size=13)
    line(personal.get("location", ""), size=10, color=(90, 90, 90))
    line(personal.get("email", ""), size=10, color=(90, 90, 90))
    line(personal.get("phone", ""), size=10, color=(90, 90, 90))
    line(personal.get("linkedin", ""), size=10, color=(90, 90, 90))

    doc.add_paragraph()  # spacer

    # ── Date ──────────────────────────────────────────────────────────────────
    line(date.today().strftime("%d %B %Y"))

    doc.add_paragraph()  # spacer

    # ── Addressee block ───────────────────────────────────────────────────────
    if job.get("title"):
        line(_clean_job_title(job["title"]), bold=True)
    if job.get("company"):
        line(job["company"])

    doc.add_paragraph()  # spacer

    # ── Salutation ────────────────────────────────────────────────────────────
    line("Dear Hiring Team,")
    doc.add_paragraph()

    # ── Body ──────────────────────────────────────────────────────────────────
    cover_text = job.get("cover_letter") or ""
    if cover_text:
        # Split on double newlines; fall back to single if needed
        paras = [p.strip() for p in cover_text.split("\n\n") if p.strip()]
        if len(paras) == 1:
            paras = [p.strip() for p in cover_text.split("\n") if p.strip()]
        for para_text in paras:
            para_text = re.sub(r"\*\*(.+?)\*\*", r"\1", para_text)  # strip markdown bold
            p = doc.add_paragraph()
            _spacing(p, before=0, after=10)
            _font(p.add_run(para_text), 11)
    else:
        p = doc.add_paragraph()
        _font(p.add_run("[Cover letter text not available.]"), 11, italic=True)

    # ── Sign-off ──────────────────────────────────────────────────────────────
    line("Yours sincerely,")
    doc.add_paragraph()  # space for physical signature
    doc.add_paragraph()
    line(personal["name"], bold=True)

    doc.save(path)
    print(f"[Documents] Cover letter saved: {path.name}")
