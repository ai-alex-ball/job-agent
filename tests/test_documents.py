"""
tests/test_documents.py — document generation unit tests

Run with: python3 -m pytest tests/ -v
"""
import json
import sys
import tempfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from documents import _build_cv, _parse_tailored_cv

# Fake profile data so this suite doesn't depend on a real, personal
# profile.json being present (profile.json is gitignored — see
# profile.example.json for the template).
PROFILE = {
    "personal": {
        "name": "Jane Doe",
        "location": "London, UK",
        "email": "jane.doe@example.com",
        "phone": "07700 900123",
        "linkedin": "https://www.linkedin.com/in/janedoe/",
    },
    "job_preferences": {
        "target_roles": ["Programme Director", "Innovation Lead"],
    },
    "summary": "Test candidate summary.",
    "key_metrics": {
        "years_experience": 15,
        "startups_mentored": 60,
        "startups_incubated": 40,
        "portfolio_raised": "£5m+",
        "portfolio_market_cap_managed": "$1B+",
    },
    "career_history": [
        {
            "company": "Acme Corp",
            "location": "London",
            "title": "Programme Director",
            "start": "2021",
            "end": "present",
            "type": "full-time",
            "achievements": ["Led 8 cohorts", "Raised £5m across portfolio"],
        }
    ],
    "ai_certifications": [
        {"issuer": "Anthropic", "title": "Claude 101", "year": 2026},
    ],
    "education": ["BSc Computer Science — Example University"],
    "skills": {
        "ai_specific": ["AI programme design", "AI product strategy"],
        "technology": ["Artificial intelligence", "Fintech"],
        "programme_leadership": ["Accelerator design", "Portfolio management"],
        "venture_and_investment": ["Venture building", "Due diligence"],
        "leadership": ["Team leadership", "Stakeholder management"],
    },
}

FIXTURE_JOB = {
    "title": "Test Role",
    "company": "Test Co",
    "location": "London",
    "matched_skills": "[]",
}


def _all_table_text(doc: Document) -> str:
    """Extract all text from all tables in the document."""
    parts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _generate_cv(job: dict = FIXTURE_JOB, profile: dict = PROFILE) -> Document:
    """Generate a CV docx and return it as a python-docx Document."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        tmp_path = Path(f.name)
    try:
        matched_skills = json.loads(job.get("matched_skills") or "[]")
        _build_cv(profile, job, matched_skills, tmp_path)
        return Document(tmp_path)
    finally:
        tmp_path.unlink(missing_ok=True)


# ---------------------------------------------------------------------------
# Career at a Glance tests
# ---------------------------------------------------------------------------

class TestCareerAtAGlance:
    def test_years_experience_appears_in_table(self):
        doc = _generate_cv()
        assert "15 Years" in _all_table_text(doc), "Missing '15 Years' in Career at a Glance"

    def test_startups_mentored_appears_in_table(self):
        doc = _generate_cv()
        assert "60+" in _all_table_text(doc), "Missing '60+' in Career at a Glance"

    def test_startups_incubated_appears_in_table(self):
        doc = _generate_cv()
        table_text = _all_table_text(doc)
        assert "40" in table_text, "Missing '40' in Career at a Glance"

    def test_portfolio_raised_appears_in_table(self):
        doc = _generate_cv()
        assert "£5m+" in _all_table_text(doc), "Missing '£5m+' in Career at a Glance"

    def test_portfolio_market_cap_appears_in_table(self):
        doc = _generate_cv()
        assert "$1B+" in _all_table_text(doc), "Missing '$1B+' in Career at a Glance"

    def test_all_five_labels_present(self):
        doc = _generate_cv()
        table_text = _all_table_text(doc)
        for label in ("Experience", "Startups Mentored", "Startups Incubated",
                      "Portfolio Raised", "Portfolio Managed"):
            assert label in table_text, f"Missing label '{label}' in Career at a Glance"


# ---------------------------------------------------------------------------
# Skills table tests
# ---------------------------------------------------------------------------

class TestSkillsTable:
    def test_ai_technology_category_present(self):
        doc = _generate_cv()
        assert "AI & Technology" in _all_table_text(doc)

    def test_programme_leadership_category_present(self):
        doc = _generate_cv()
        assert "Programme Leadership" in _all_table_text(doc)

    def test_skills_use_dot_separator(self):
        doc = _generate_cv()
        table_text = _all_table_text(doc)
        # Skills should use · not commas
        assert "·" in table_text, "Expected · separator in skills table"


# ---------------------------------------------------------------------------
# _parse_tailored_cv tests
# ---------------------------------------------------------------------------

class TestParseTailoredCv:
    def test_returns_none_for_empty(self):
        assert _parse_tailored_cv(None) is None
        assert _parse_tailored_cv("") is None

    def test_returns_none_for_plain_text(self):
        assert _parse_tailored_cv("Jane Doe — Senior Manager\n\nExperience\nAcme Corp...") is None

    def test_parses_valid_json(self):
        data = {
            "summary": "Test summary",
            "roles": [{"company": "Acme Corp", "title": "Director", "bullets": ["Led 8 cohorts"]}],
            "projects": []
        }
        result = _parse_tailored_cv(json.dumps(data))
        assert result is not None
        assert result["roles"][0]["company"] == "Acme Corp"

    def test_strips_code_fences(self):
        data = {"summary": "x", "roles": [], "projects": []}
        fenced = f"```json\n{json.dumps(data)}\n```"
        result = _parse_tailored_cv(fenced)
        assert result is not None

    def test_returns_none_if_no_roles_key(self):
        assert _parse_tailored_cv('{"summary": "x"}') is None


# ---------------------------------------------------------------------------
# Structured tailored_cv integration test
# ---------------------------------------------------------------------------

class TestStructuredCvBullets:
    def test_uses_structured_bullets_when_present(self):
        structured = {
            "summary": "Test summary",
            "roles": [
                {
                    "company": "Acme Corp",
                    "title": "Global Programme Director",
                    "bullets": ["STRUCTURED BULLET unique marker 12345"]
                }
            ],
            "projects": []
        }
        job = {**FIXTURE_JOB, "tailored_cv": json.dumps(structured)}
        doc = _generate_cv(job)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "STRUCTURED BULLET unique marker 12345" in full_text

    def test_renders_projects_section_when_present(self):
        structured = {
            "summary": "Test summary",
            "roles": [],
            "projects": [
                {
                    "name": "AI Job Application Agent",
                    "date": "March 2026",
                    "bullets": ["Built agentic system using Claude API"]
                }
            ]
        }
        job = {**FIXTURE_JOB, "tailored_cv": json.dumps(structured)}
        doc = _generate_cv(job)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # _section_heading uppercases the text
        assert "RECENT PROJECTS" in full_text
        assert "AI Job Application Agent" in full_text

    def test_falls_back_to_profile_bullets_when_no_match(self):
        structured = {
            "summary": "Test summary",
            "roles": [
                {"company": "SomeOtherCompany", "title": "Director", "bullets": ["Bullet A"]}
            ],
            "projects": []
        }
        job = {**FIXTURE_JOB, "tailored_cv": json.dumps(structured)}
        doc = _generate_cv(job)
        # Company names are in table cells (job title rows); combine all text
        all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + _all_table_text(doc)
        # Acme Corp company name should appear (fallback to profile roles)
        assert "Acme Corp" in all_text
